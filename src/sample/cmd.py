import time
import os
import xml.dom.minidom

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor

# pyinstaller -F -p E:\project\py\python-example\venv\Lib\site-packages;E:\project\py\python-example\venv\Lib\site-packages\lxml cmd.py

def create_book_cover(document,bookName,authorName,public):
    document.add_paragraph("\n\n\n\n")
    # 书名
    bookNameHeader = document.add_paragraph()
    book_name_run = bookNameHeader.add_run(bookName)
    # 居中对齐
    book_name_header_format = bookNameHeader.paragraph_format
    book_name_header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 字体颜色
    book_name_run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小，22 == 二号
    book_name_run.font.size = Pt(22)
    # 加粗
    book_name_run.bold = True

    document.add_paragraph("\n\n\n\n")
    # 作者
    auther_header = document.add_paragraph()
    auther_run = auther_header.add_run(authorName)
    # 居中对齐
    auther_header_format = auther_header.paragraph_format
    auther_header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 字体颜色
    auther_run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小
    auther_run.font.size = Pt(14)

    #  版权声明
    public_header = document.add_paragraph()
    public_run = public_header.add_run("\n\n\n\n\n"+public)
    # 左对齐
    public_header_format = public_header.paragraph_format
    public_header_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 字体颜色
    public_run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小
    public_run.font.size = Pt(14)


def create_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    # 两端对齐
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # 首行缩进，没找到字符
    paragraph_format.first_line_indent = Cm(0.74)

    # 字体大小，14 == 四号
    run.font.size = Pt(14)
    # run.bold = True
    # 设置字体
    run.font.name = 'Calibri'
    #     行间距
    paragraph_format.line_spacing = Pt(22)

    #  段落间距 /段前
    # paragraph_format.space_before = Pt(22)


def create_head(document, title):
    header = document.add_heading()
    run = header.add_run(title)
    # 居中对齐
    header_format = header.paragraph_format
    header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 字体颜色
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小，18 == 小二
    run.font.size = Pt(18)
    # 加粗
    run.bold = True


def parse_xml(path):
    dom = xml.dom.minidom.parse(path)
    root = dom.documentElement
    book_name_node = root.getElementsByTagName('name')[0]
    author_node = root.getElementsByTagName('author')[0]
    book_name = book_name_node.childNodes[0].data
    author = author_node.childNodes[0].data
    return book_name.strip(), author.strip()


def read_dir(dir):
    L = []
    M = []
    path_list = os.listdir(dir)
    for file in path_list:
        if os.path.splitext(file)[1] == '.txt':
            L.append(file)
    L.sort(key=lambda x: int(x.split('.')[0]))
    for file in L:
        M.append(os.path.join(dir, file))
    return M, os.path.join(dir, "book.xml"), os.path.join(dir, "chapter.xml")


def separation(ary):
    L = []
    i = 0
    length = len(ary)
    if length < 300:
        L.append(ary)
        return L
    while length > 250:
        L.append(ary[i:i + 150])
        i += 150
        length -= 150
    L.append(ary[i::])
    return L


def mkdir(path):
    print("创建目录[ " + path+" ]")
    path = path.strip()
    # 去除尾部 \ 符号
    path = path.rstrip("\\")
    # 判断路径是否存在
    isExists = os.path.exists(path)
    # 判断结果
    if not isExists:
        # 如果不存在则创建目录,创建目录操作函数
        '''
        os.mkdir(path)与os.makedirs(path)的区别是,当父目录不存在的时候os.mkdir(path)不会创建，os.makedirs(path)则会创建父目录
        '''
        # 此处路径最好使用utf-8解码，否则在磁盘中可能会出现乱码的情况
        os.makedirs(path)
        print("目录[ "+ path + ' ]创建成功')
        return True
    else:
        # 如果目录存在则不创建，并提示目录已存在
        print("目录[ " + path + ' ]已存在')
        return False


if __name__ == '__main__':
    print("开始执行转换 ...")
    # 获取当前目录
    _dir = os.getcwd()
    # _dir = "E:\\book_en\\23_2073"
    print("当前目录：" + _dir)
    print("正在读取文件...")
    ary, book_xml_file_path, chapter_xml_file_path = read_dir(_dir)
    txt_file_len = len(ary)
    print("txt 文件数量：" + str(txt_file_len))
    print("读取文件：book.xml")
    bookname, author = parse_xml(book_xml_file_path)
    print("书籍 ：" + bookname + "  作者 ：" + author)
    print("正在进行分章...")
    list = separation(ary)
    dist_dir = _dir + "/dist/"
    mkdir(dist_dir)
    book_name = "book"
    book_index = 1

    for book in list:
        print("长度：" + str(len(book)) + "  start:" + book[0] + "  end:" + book[-1])
        document = Document()
        create_book_cover(document, bookname, author, "Published By Sofanovel")
        document.add_page_break()
        this_book_name = book_name + "_" + str(book_index)+".docx"

        for seq in book:
            first = True
            for line2 in open(seq, encoding="utf-8"):
                if first:
                    # 第一行是标题
                    create_head(document, line2)
                else:
                    text = line2
                    create_paragraph(document, text.strip())

                first = False

            document.add_page_break()
        document.save(dist_dir+this_book_name)
        print("文件[ " + this_book_name + " ]生成完成...")
        book_index += 1

    print("执行完成 ...")
    os.startfile(dist_dir)
