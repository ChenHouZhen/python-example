import os
import xml.dom.minidom

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor


def create_book_cover(document,bookName,authorName,public):
    document.add_paragraph("\n\n\n\n")
    # 书名
    bookNameHeader = document.add_paragraph()
    book_name_run = bookNameHeader.add_run(bookName)
    # 居中对齐
    book_name_header_format = bookNameHeader.paragraph_format
    book_name_header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 字体颜色
    book_name_run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小，22 == 二号
    book_name_run.font.size = Pt(22)
    # 加粗
    book_name_run.bold = True

    document.add_paragraph("\n\n\n\n")
    # 作者
    auther_header = document.add_paragraph()
    auther_run = auther_header.add_run(authorName)
    # 居中对齐
    auther_header_format = auther_header.paragraph_format
    auther_header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 字体颜色
    auther_run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小
    auther_run.font.size = Pt(14)

    #  版权声明
    public_header = document.add_paragraph()
    public_run = public_header.add_run("\n\n\n\n\n"+public)
    # 左对齐
    public_header_format = public_header.paragraph_format
    public_header_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 字体颜色
    public_run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小
    public_run.font.size = Pt(14)


def create_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    # 两端对齐
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # 首行缩进，没找到字符
    paragraph_format.first_line_indent = Cm(0.74)

    # 字体大小，14 == 四号
    run.font.size = Pt(14)
    # run.bold = True
    # 设置字体
    run.font.name = 'Calibri'
    #     行间距
    paragraph_format.line_spacing = Pt(22)

    #  段落间距 /段前
    # paragraph_format.space_before = Pt(22)


def create_head(document, title):
    header = document.add_heading()
    run = header.add_run(title)
    # 居中对齐
    header_format = header.paragraph_format
    header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 字体颜色
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 字体大小，18 == 小二
    run.font.size = Pt(18)
    # 加粗
    run.bold = True


def parse_xml(path):
    dom = xml.dom.minidom.parse(path)
    root = dom.documentElement
    booknameNode = root.getElementsByTagName('name')[0]
    authorNode = root.getElementsByTagName('author')[0]

    print(booknameNode.childNodes[0].data)
    print(authorNode.childNodes[0].data)
    bookname = booknameNode.childNodes[0].data
    author = authorNode.childNodes[0].data
    return bookname.strip(), author.strip()


def read_dir(dir):
    L = []
    M = []
    path_list = os.listdir(dir)
    for file in path_list:
        if os.path.splitext(file)[1] == '.txt':
            L.append(file)
    L.sort(key=lambda x: int(x.split('.')[0]))
    print(L)
    for file in L:
        M.append(os.path.join(dir, file))

    print(M)
    return M, os.path.join(dir, "book.xml"), os.path.join(dir, "chapter.xml")


def separation(ary):
    print("长度：len="+ str(len(ary)))
    L = []
    i = 0
    length = len(ary)
    if length < 300:
        L.append(ary)
        return L
    while length > 250:
        L.append(ary[i:i + 150])
        i += 150
        length -= 150
    L.append(ary[i::])
    return L

if __name__ == '__main__':

    ary, book_path, chapter_path = read_dir("E:\\book_en\\23_3487")
    bookname, author = parse_xml(book_path)
    list = separation(ary)
    book_name = "book"
    book_index = 1
    for book in list:
        print("长度：" + str(len(book)) + "  start:" + book[0] + "  end:" + book[-1])
        document = Document()
        create_book_cover(document, bookname, author, "Published By Sofanovel")
        document.add_page_break()

        for seq in book:
            first = True
            for line2 in open(seq, encoding="utf-8"):
                if first:
                    # 第一行是标题
                    create_head(document, line2)
                else:
                    text = line2
                    create_paragraph(document, text.strip())

                first = False

            document.add_page_break()
        print(book_path)
        document.save(book_name+"_" + str(book_index)+".docx")
        book_index += 1